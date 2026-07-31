"""启文技术交底书 v2.0.0 — 含AI Agent系统发明"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs: run.font.name = 'Microsoft YaHei'
    return p
def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs: run.font.name = 'Microsoft YaHei'
    return p
def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs: run.font.name = 'Microsoft YaHei'
    return p
def para(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10.5)
    return p
def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs: run.font.name = 'Microsoft YaHei'; run.font.size = Pt(10)
    return p
def tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Microsoft YaHei'; run.font.size = Pt(9); run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = cell._element.get_or_add_tcPr()
        shd = shading_elm.makeelement(qn('w:shd'), {qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear'})
        shading_elm.append(shd)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val or ''))
            run.font.name = 'Microsoft YaHei'; run.font.size = Pt(9)
    return table
def page_break(): doc.add_page_break()

# ── Cover ──
for _ in range(5): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('技 术 交 底 书'); r.font.name = 'Microsoft YaHei'; r.font.size = Pt(26); r.bold = True
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('（计算机软件发明专利申请）'); r.font.name = 'Microsoft YaHei'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
for _ in range(4): doc.add_paragraph()
tbl(['项目', '内容'],
    [['发明名称', '一种AI原生的跨编辑器智能写作Agent系统及方法'],
     ['发明人', '周铭翀'],
     ['所属单位', 'BitWool'],
     ['申请日期', '2026年7月'],
     ['技术类别', '计算机软件 / 人工智能 / 自然语言处理 / 知识管理'],
     ['联系方式', 'bitwool@163.com']])

page_break()
h1('一、技术领域')
para('本发明属于计算机软件技术领域，具体涉及桌面端应用软件中的人工智能写作Agent系统，包括结构化标签驱动的跨编辑器操作机制、多编辑器统一桥接框架、面板网格可拖拽布局系统，以及基于对抗验证的AI输出可靠性保障机制。')

h1('二、现有技术介绍及存在的不足')
h2('（一）现有技术情况简介')
para('当前AI辅助写作技术主要分为三类：第一类为云端对话式AI（如ChatGPT），以会话界面为主，缺乏对写作编辑器内容的直接操控能力；第二类为编辑器内嵌AI（如Notion AI、腾讯文档AI），仅支持文内插入生成内容，不支持跨文档类型的统一操作调度；第三类为桌面Agent型AI（如Open Interpreter），可执行系统级命令但缺乏对专业编辑器（文档/PPT/白板/思维导图）的结构化操作抽象。')

h2('（二）现有技术方案存在的技术问题')
para('上述现有技术方案存在以下主要不足：')
bullet('（1）缺乏跨编辑器统一操作机制：现有AI写作工具仅能操作单一类型编辑器（通常是文档编辑器），无法统一调度PPT幻灯片、白板画布、思维导图节点等异构编辑器。用户需要在不同AI助手间切换，工作流严重割裂。')
bullet('（2）AI输出缺乏结构化标签和分级安全控制：现有AI写作工具将生成内容以纯文本形式输出，缺乏对"追加/插入/替换/重写/删除"等操作类型的结构标记，更缺乏对操作安全性的分级控制（安全操作自动执行，危险操作需确认）。')
bullet('（3）基于Web技术的写作软件存在窗口关闭导致内容丢失的共性问题：window.beforeunload事件无法执行异步Promise操作，导致用户关闭窗口时正在执行的写入被中断。')
bullet('（4）插件系统依赖网络下载，在政务网、内网等受限场景功能受限。')
bullet('（5）缺乏对AI输出质量的对抗验证机制，AI生成内容可能包含事实性错误未被检测。')

h2('（三）本发明解决思路')
para('为解决上述技术问题，本发明提出了一种AI原生的跨编辑器智能写作Agent系统，通过XML Action标签驱动的编辑器操作机制、四编辑器统一桥接框架（EditorBridge）、主进程-渲染进程双向握手保存协议、全量离线插件打包、以及对抗验证模型等创新，实现了AI对多种编辑器内容的直接、安全、可靠操控。')

page_break()
h1('三、发明内容与核心技术方案')
h2('（一）系统总体架构')
para('本发明采用Electron跨平台桌面应用框架构建，创新性地引入了Panel Grid面板网格系统（取代传统固定侧边栏）和AI Agent引擎作为独立应用层：')

tbl(['架构层次', '主要组成模块', '核心功能'],
    [['表现层', 'React 18 + TypeScript + Redux Toolkit + framer-motion', '用户界面、编辑器渲染、Panel Grid面板管理、动画交互'],
     ['应用层', 'AI Agent引擎（ChatPanel）、Action解析器/执行器、EditorBridge桥接器', 'Agent对话、XML标签解析、跨编辑器操作调度'],
     ['面板层', 'PanelGrid / SplitPane / FloatingPanel / DragOverlay', '可拖拽分屏、浮动窗口、布局持久化'],
     ['桥接层', 'contextBridge + 40+ IPC安全通道', '渲染进程-主进程安全通信路由'],
     ['主进程层', 'Electron 31 + Node.js + better-sqlite3', '窗口管理、文件系统、数据库、HTTP代理'],
     ['存储层', 'better-sqlite3 + Supabase PostgreSQL', '本地持久化 + 可选云同步'],
     ['协作层', 'Y.js CRDT + Supabase Realtime', '多光标实时协作、冲突自动合并']])

h2('（二）技术特征一：XML Action标签驱动的跨编辑器智能操作机制【核心创新点一】')
para('本发明最核心的创新在于设计了一套基于XML结构化标签的AI编辑器操作语言，使AI Agent可以精确操控四种异构编辑器（文档、PPT幻灯片、白板画布、思维导图）。')
h3('1. Action标签语法')
para('AI在流式SSE回复中嵌入<action> XML标签，由actionParser.ts模块中的正则表达式实时解析。标签格式如下：')
para('<action type="append|insert|replace|rewrite|delete" target="document|slides|whiteboard|mindmap" position="cursor|start|end" section="可选">操作内容</action>')
para('其中type指定操作类型，target指定目标编辑器，position指定插入位置，section用于精确匹配待替换的段落内容。')
h3('2. 分级安全控制机制')
para('本发明将5种操作按安全等级分为两类：append和insert为"安全操作"，AI自动执行无需用户确认，执行结果直接写入编辑器并触发绿色闪烁视觉反馈；replace、rewrite和delete为"需确认操作"，执行前弹出ActionConfirm组件展示Diff对比卡片（before/after），用户通过approveDiff或rejectDiff函数批准或拒绝操作。拒绝后显示红色短暂提示并恢复原内容。')
h3('3. 兜底机制')
para('AI回复中未使用<action>标签的内容分两种情况处理：短于80字的内容直接以对话气泡形式显示；超过80字的内容自动作为append操作插入当前编辑器光标位置，确保大量生成内容不丢失。同时<plan>标签渲染为任务计划卡片，<thinking>标签渲染为打字机逐字动画的可折叠思考过程块。')

h2('（三）技术特征二：四编辑器统一桥接框架（EditorBridge）【核心创新点二】')
para('本发明设计了一个基于注册表模式（Map<EditorType, EditorAPI>）的多编辑器统一桥接框架：')
bullet('注册机制：每种编辑器在挂载时调用editorBridge.register(type, api)注册自身的EditorAPI实现（包含insertText、replaceSection、deleteSection、getContent、getSelection等方法）。')
bullet('动态路由：AI Agent通过editorBridge.getActive()获取当前活跃编辑器的桥接接口，无需感知底层编辑器的具体实现差异。')
bullet('操作反馈：通过executeWithFeedback(editorType, fn)统一处理操作后的视觉反馈——含状态横幅（成功/失败/需确认）、编辑器内容闪烁、操作日志记录。')
bullet('四编辑器状态：文档编辑器（TipTap）完全接入，PPT/Slides、白板、思维导图均已注册桥接接口，AI可统一调度操作。')

h2('（四）技术特征三：Panel Grid可拖拽面板布局系统【核心创新点三】')
para('本发明将传统桌面应用的固定侧边栏布局替换为可自由拖拽的Panel Grid面板网格系统：')
bullet('VS Code风格拖拽分屏：拖拽面板标题栏至屏幕边缘32px热区即触发分屏，支持水平和垂直方向。')
bullet('独立浮动窗口：面板可从网格中脱离为独立浮动窗口（iPad台前调度效果），自由移动和调整大小。浮动窗口拖至屏幕边缘44px范围内自动吸附回网格。')
bullet('布局持久化：每个文档独立记住其面板布局配置（panel_layouts表），重启后自动恢复上次布局。')

page_break()
h2('（五）技术特征四：双向握手保存协议【核心创新点四】')
para('针对Web技术栈写作工具"关闭窗口导致内容丢失"的共性问题，本发明设计了基于Electron IPC的主进程-渲染进程双向握手保存协议，分为四个阶段：')
bullet('拦截阶段：主进程监听close事件，调用e.preventDefault()阻止窗口立即关闭。')
bullet('通知阶段：主进程通过webContents.send向渲染进程发送app-before-close信号。')
bullet('写入阶段：渲染进程收到信号后执行await autoSave.flushAll()，等待所有pending文档的异步IPC写入完成。')
bullet('确认阶段：渲染进程通过flush-complete信号通知主进程保存完成，主进程设置isReallyClosing=true后执行真正的窗口关闭。同时设3秒超时兜底。')

h2('（六）技术特征五：全量离线插件打包与职业化推荐【核心创新点五】')
para('本发明将所有插件代码在编译阶段完整打包于安装包中（pluginRegistry.ts静态注册表），实现了零网络依赖的全功能可用。内置PROFESSION_PLUGIN_MAP注册表定义了6类职业（学术研究/法律工作/教育培训/医疗健康/内容创作/通用知识）与17+插件之间的对应关系。首次启动时用户在引导页选择职业类型，系统自动调用getPluginsForProfession(profession)检索并激活对应插件组合。')

h2('（七）技术特征六：RDM科研数据管理扩展模块【核心创新点六】')
para('本发明在插件系统中集成了一套面向实验室科研场景的重量级扩展插件RDM（Research Data Management），包含8个子模块（仪表盘/项目管理/电子实验记录ELN/库存管理/仪器管理/审计追踪/审批流程/报告生成），具备独立的better-sqlite3数据库Schema、Redux状态管理Slice和完整UI组件。同时提供用户自定义插件SDK（sdk/UserPluginEditor.tsx），基于Monaco Editor代码高亮和iframe沙箱隔离，支持开发者按需扩展组织专属功能。')

page_break()
h1('四、具体实施方式')
h2('方案一（主方案）：AI Agent跨编辑器协同写作')
para('以学术研究者用户为例，完整实施方式如下：')
bullet('（1）用户下载并安装启文桌面端（Windows/macOS/Linux），首次启动触发引导页，选择"学术研究"职业，系统自动创建工作区并激活引用格式生成、论文大纲助手、关键词提取三个核心插件。')
bullet('（2）用户在文档编辑器中撰写论文初稿，通过快捷键唤起ChatPanel对话面板（Panel Grid分屏布局，编辑器60%+ChatPanel 40%）。')
bullet('（3）用户在对话面板中输入"帮我检查引言部分并补充相关文献综述"，AI Agent进入"分析需求→构思方案→生成内容→检查润色"状态轮播后，通过流式SSE输出结构化回复，其中包含<action type="insert" target="document" position="cursor">标签，actionExecutor自动将生成的文献综述段落插入编辑器光标位置。')
bullet('（4）若AI发现引言中某段表述不准确需要改写，回复中包含<action type="rewrite" target="document" section="原段落内容" old="旧表述">新表述</action>，ActionConfirm弹出Diff卡片等待用户确认。')
bullet('（5）用户关闭窗口时，双向握手协议确保所有编辑内容和对话历史完整写入better-sqlite3数据库。下次打开时，文档内容和对话上下文完整恢复。')

h2('方案二（优化方案）：RDM科研全流程管理')
para('在方案一基础上，安装RDM科研数据管理插件，实现从文献调研（doiApi查询）→ 实验记录（ELN电子实验记录）→ 数据分析（仪表盘统计）→ 论文写作（AI Agent辅助）→ 发表提交（Word/PDF/LaTeX导出）的全流程闭环。审批流程（Approvals）和审计追踪（Audit）模块满足GLP合规要求。')

h2('方案三（云端协作方案）')
para('在方案一基础上，注册BitWool账号并开启云同步。通过Y.js CRDT技术实现同一文档的多光标实时协作，各协作者的编辑操作通过Supabase Realtime通道广播，CRDT算法保证离线编辑重新联网后的无冲突自动合并。')

page_break()
h1('五、有益的技术效果')
bullet('（1）跨编辑器统一操作效率：AI Agent通过单一对话界面即可操作文档、PPT、白板、思维导图四种编辑器，消除用户在多个AI工具间的切换成本。')
bullet('（2）操作安全可控性：分级安全机制确保安全操作（append/insert）自动执行，危险操作（replace/rewrite/delete）经用户确认，AI始终不越过用户授权边界。')
bullet('（3）数据可靠性：双向握手协议从根本上解决了窗口关闭导致内容丢失的问题，最多仅丢失600ms内的写作内容。')
bullet('（4）离线完整性：全量插件打包和本地数据库确保无网络环境下100%核心功能可用，适用于内网、政务网等受限场景。')
bullet('（5）布局灵活性：Panel Grid系统允许用户自由拖拽面板和创建浮动窗口，相比固定侧边栏布局的同类产品，空间利用率和个性化程度显著提升。')
bullet('（6）科研场景覆盖度：RDM插件的8个子模块覆盖了实验室科研管理的完整流程，填补了通用写作工具在科研垂直场景中的空白。')

output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '启文_技术交底书_v2.0.0.docx')
doc.save(output_path)
print('Written:', output_path)

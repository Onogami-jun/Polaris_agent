"""
启文 程序鉴别材料 v2.0.0 生成器
中国软件著作权登记 — 源程序鉴别材料
前30页 + 后30页
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, glob, re

# ── Files to include ──
QIWEN_ROOT = r'D:\github\qiwen_v2.0.0'

# Key source files for first 30 pages (core modules)
FIRST_FILES = [
    'src/renderer/panels/ChatPanel.tsx',
    'src/renderer/panels/actionParser.ts',
    'src/renderer/panels/actionExecutor.ts',
    'src/renderer/panels/editorBridge.ts',
    'src/renderer/panels/types.ts',
    'src/renderer/panels/AgentControlBar.tsx',
    'src/renderer/panels/ActionConfirm.tsx',
    'electron/main.js',
    'electron/preload.js',
    'src/renderer/App.tsx',
    'src/renderer/components/editor/MarkdownEditor.tsx',
    'src/renderer/components/editor/EditorToolbar.tsx',
    'src/renderer/plugins/pluginRegistry.ts',
    'src/renderer/plugins/PluginSidebarPanel.tsx',
    'src/renderer/store/index.ts',
    'src/renderer/store/slices/appSlice.ts',
    'src/renderer/store/slices/documentsSlice.ts',
]

# Key source files for last 30 pages (services + plugins + rdm)
LAST_FILES = [
    'src/renderer/services/syncEngine.ts',
    'src/renderer/services/cloudSync.ts',
    'src/renderer/plugins/rdm/RdmPlugin.tsx',
    'src/renderer/plugins/rdm/db/schema.ts',
    'src/renderer/plugins/rdm/db/operations.ts',
    'src/renderer/plugins/rdm/store/slices/rdmSlice.ts',
    'src/renderer/plugins/api/doiApi.ts',
    'src/renderer/plugins/api/semanticApi.ts',
    'src/renderer/plugins/api/readabilityApi.ts',
    'src/renderer/plugins/api/legalApi.ts',
    'src/renderer/plugins/api/icdApi.ts',
    'src/renderer/plugins/api/drugApi.ts',
    'src/renderer/utils/autoSave.ts',
    'src/renderer/utils/ipc.ts',
    'src/renderer/lib/supabase.ts',
    'src/renderer/components/stats/WritingStatsView.tsx',
    'src/renderer/components/stats/DocumentGraphView.tsx',
    'src/renderer/components/code/CodeViewer.tsx',
    'src/renderer/i18n/index.ts',
]

def read_file(relpath):
    """Read file content, return (filename, lines)"""
    fullpath = os.path.join(QIWEN_ROOT, relpath)
    if not os.path.exists(fullpath):
        return (os.path.basename(relpath), [f'// FILE NOT FOUND: {relpath}'])
    with open(fullpath, 'r', encoding='utf-8', errors='replace') as f:
        return (os.path.basename(relpath), f.readlines())

def create_doc(files, doc_title, output_name):
    doc = Document()
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(8)
    style.paragraph_format.line_spacing = 1.1
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── Cover ──
    for _ in range(4): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(doc_title); r.font.name = 'Microsoft YaHei'; r.font.size = Pt(22); r.bold = True
    doc.add_paragraph()
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('启文 QiWen Writer v2.0.0'); r.font.name = 'Microsoft YaHei'; r.font.size = Pt(14)
    doc.add_paragraph()
    info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run('BitWool 数字毛织 · bitwool.cn\n2026年7月'); r.font.name = 'Microsoft YaHei'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x88,0x88,0x88)

    doc.add_page_break()

    # ── Code listing ──
    line_no = 1
    for relpath in files:
        filename, lines = read_file(relpath)

        # Section header
        p = doc.add_paragraph()
        run = p.add_run(f'┌─ 文件: {relpath}  ({len(lines)} 行) ─┐')
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(7)
        run.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        for l in lines:
            l = l.rstrip('\n\r')
            if l.strip() == '':
                p = doc.add_paragraph(' ')
            else:
                p = doc.add_paragraph()
                run = p.add_run(l[:120])
                run.font.name = 'Consolas'
                run.font.size = Pt(7)

            line_no += 1

        # File separator
        p = doc.add_paragraph('─' * 80)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), output_name)
    doc.save(output_path)
    print(f'Written: {output_path}  ({len(files)} files)')

# Generate
create_doc(FIRST_FILES, '源程序鉴别材料（前30页）', '启文_程序鉴别材料_前30页.docx')
create_doc(LAST_FILES, '源程序鉴别材料（后30页）', '启文_程序鉴别材料_后30页.docx')

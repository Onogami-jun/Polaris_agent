"""启文技术文档 v2.0.0 生成器 — Python版"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
sections = doc.sections
for section in sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style helpers ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
    return p

def para(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    if bold_parts:
        for part, is_bold in bold_parts:
            run = p.add_run(part)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10.5)
            run.bold = is_bold
    else:
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10.5)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10)
    return p

def tbl(headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(9)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Shade header
        shading_elm = cell._element.get_or_add_tcPr()
        shd = shading_elm.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F2F2F2',
            qn('w:val'): 'clear',
        })
        shading_elm.append(shd)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val or ''))
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    return table


def page_break():
    doc.add_page_break()


def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('启  文')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(28)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('QiWen Studio')
run.font.name = 'Consolas'
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

motto = doc.add_paragraph()
motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = motto.add_run('启于思，行于文')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

separator()

for _ in range(2):
    doc.add_paragraph()

doctitle = doc.add_paragraph()
doctitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = doctitle.add_run('技 术 文 档')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(22)
run.bold = True

for _ in range(8):
    doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('v2.0.0 · 2026年7月')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('BitWool 数字毛织')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(11)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run('bitwool.cn · bitwool@163.com')
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

page_break()
# ═══════════════════════════════════════════
# 一、产品概述
# ═══════════════════════════════════════════
h1('一、产品概述')

para('启文（QiWen Studio / QiWen Writer）是由 BitWool 数字毛织自主研发的本地优先、AI 原生的专业级桌面写作与知识管理平台。启文 v2.0 将人工智能从辅助工具升级为创作协作者，通过 AI Agent 系统实现对编辑器内容的直接操控，标志着产品从"带 AI 功能的写作软件"向"AI 主导的对话式写作应用"的范式转变。',
     [('启文（QiWen Studio / QiWen Writer）是由 BitWool 数字毛织自主研发的本地优先、', False),
      ('AI 原生', True),
      ('的专业级桌面写作与知识管理平台。启文 v2.0 将人工智能从辅助工具升级为创作协作者，通过 AI Agent 系统实现对编辑器内容的直接操控，标志着产品从"带 AI 功能的写作软件"向"AI 主导的对话式写作应用"的范式转变。', False)])

h2('1.1 核心设计原则')
bullet('本地优先：所有文档默认存储于用户设备的嵌入式数据库，无需联网即可使用 100% 核心功能。')
bullet('AI 原生：AI Agent 深度集成于编辑流程，通过结构化标签直接操作编辑器内容，而非简单叠加对话窗口。')
bullet('跨文档统一：提供文档、演示文稿（PPT）、白板、思维导图四种编辑器，通过统一的桥接接口使 AI Agent 可操作任意编辑器。')
bullet('离线完整：所有插件和 AI 基础能力在编译时打包进安装包，无网络环境下功能完整可用。')
bullet('数据可靠：本地 SQLite 存储 + 关闭前握手保存协议 + 多层级自动保存，确保数据不丢失。')

h2('1.2 版本信息')
tbl(['属性', '说明'], [
    ['当前版本', 'v2.0.0'],
    ['发布日期', '2026 年 7 月'],
    ['支持平台', 'Windows 10+、macOS 12+、Linux (AppImage/deb)'],
    ['安装包格式', 'Windows: .exe NSIS 安装程序；macOS: .dmg；Linux: .deb / .AppImage'],
    ['官方网站', 'bitwool.cn'],
    ['技术栈', 'Electron 31 + React 18 + TypeScript + better-sqlite3 + TipTap'],
    ['AI 引擎', '豆包 (Doubao Seed 2.0 Pro)，火山引擎 Ark API，流式 SSE'],
    ['云服务', 'Supabase，11 表 + RLS + Realtime'],
])

page_break()
# ═══════════════════════════════════════════
# 二、系统架构
# ═══════════════════════════════════════════
h1('二、系统架构')

h2('2.1 整体架构层次')
tbl(['层次', '主要技术', '职责'], [
    ['表现层', 'React 18 + TypeScript + Redux Toolkit + framer-motion', '用户界面、编辑器渲染、面板管理、动画交互'],
    ['应用层', 'AI Agent 引擎、Action 解析/执行器、EditorBridge 桥接', 'Agent 对话、XML 标签解析、跨编辑器操作调度'],
    ['面板层', 'Panel Grid 面板网格系统', '可拖拽分屏布局、浮动窗口、独立布局记忆'],
    ['桥接层', 'contextBridge + IPC 安全通道', '渲染进程与主进程之间的安全通信路由'],
    ['主进程层', 'Electron 31 + Node.js', '窗口管理、文件系统、数据库引擎、HTTP 代理'],
    ['存储层', 'better-sqlite3 嵌入式数据库 + Supabase 云同步', '文档、设置、版本历史、插件状态的本地持久与云同步'],
    ['协作层', 'Y.js CRDT + Supabase Realtime', '多光标实时协作、冲突自动合并'],
])

h2('2.2 进程模型')
para('启文采用 Electron 标准双进程架构，严格遵循最小权限原则：')
para('主进程（Main Process）：负责窗口生命周期管理、本地文件系统访问、better-sqlite3 数据库操作、原生菜单与快捷键注册、系统对话框调用、AI 请求代理（HTTP/HTTPS 转发至火山引擎 API）、自动更新检查等所有需要系统权限的操作。')
para('渲染进程（Renderer Process）：运行在安全沙箱中（nodeIntegration: false, contextIsolation: true），负责 UI 渲染与用户交互。所有需要系统资源的操作通过 preload.js 暴露的 window.electronAPI 接口向主进程发起 IPC 调用。')
para('预加载脚本（preload.js）：通过 contextBridge.exposeInMainWorld 安全暴露 40+ 个 IPC API，涵盖文档 CRUD、工作区管理、AI 对话（流式/非流式双模）、桌面自动化、MCP 插件管理、健康检查、云同步等功能。')

h2('2.3 Panel Grid 面板系统')
para('启文 v2.0 引入了全新的 Panel Grid 面板布局系统，取代了传统的固定侧边栏布局：')
bullet('VS Code 风格拖拽分屏：拖拽面板标题栏至目标区域边缘 32px 热区即触发分屏，支持水平/垂直方向。')
bullet('浮动窗口模式：对话面板可从网格中脱离为独立的浮动窗口（iPad 台前调度效果），可自由移动和调整大小。浮动窗口拖至屏幕边缘 44px 范围内自动吸附回网格。')
bullet('布局持久化：每个文档独立记住其面板布局配置，存储于 SQLite panel_layouts 表，重启后自动恢复。')
bullet('默认布局：左侧编辑器区域（60%）+ 右侧 ChatPanel 对话面板（40%），分隔条可自由调整。')

page_break()
# ═══════════════════════════════════════════
# 三、核心功能模块
# ═══════════════════════════════════════════
h1('三、核心功能模块')

h2('3.1 AI Agent 系统（v8）')
para('启文 v2.0 的 AI Agent 系统实现了从"工具辅助"到"协同创作"的范式跃迁。Agent 不再只是生成文本，而是通过结构化 XML 标签系统直接操控编辑器内容。')

h3('3.1.1 Action 标签系统')
para('AI 的回复中可嵌入 <action> XML 标签，实现对编辑器内容的精确操作。标签由 actionParser.ts 解析，actionExecutor.ts 执行，支持以下 5 种操作类型：')

tbl(['Action 类型', '目标编辑器', '安全级别', '说明'], [
    ['append', 'document/slides/whiteboard/mindmap', '安全（自动执行）', '追加内容到文档末尾或指定位置'],
    ['insert', 'document/slides/whiteboard/mindmap', '安全（自动执行）', '在光标位置或指定位置插入内容'],
    ['replace', '任意编辑器', '需确认（Diff 卡片）', '替换指定段落，展示 before/after 对比'],
    ['rewrite', '任意编辑器', '需确认（Diff 卡片）', '重写指定内容，保留核心信息'],
    ['delete', '任意编辑器', '需确认', '删除指定内容段落'],
])

para('安全分级机制：append 和 insert 属于"安全操作"，Agent 自动执行无需用户确认，结果直接写入编辑器并显示绿色闪烁反馈。replace、rewrite、delete 属于"需确认操作"，执行前弹出 ActionConfirm 组件展示 Diff 对比卡片，用户可批准或拒绝。AI 回复中未使用 <action> 标签的内容直接显示在对话气泡中。<plan> 标签渲染为任务计划卡片，<thinking> 标签渲染为可折叠的思考过程块。')

h3('3.1.2 对话面板（ChatPanel）')
para('ChatPanel 是 AI Agent 的核心交互界面，与编辑器通过 Panel Grid 分屏常驻显示。主要特性包括：')
bullet('流式 SSE 渲染：支持标准流式和非流式双模，发生错误时自动降级到非流式模式。')
bullet('思考过程可视化：AI 回答中的推理过程以折叠块形式展示，含逐字打印动画和 5 秒后自动折叠。')
bullet('状态轮播动画：依次显示"分析需求→构思方案→生成内容→检查润色"四种状态。')
bullet('自动模式：用户发送"自动模式"后，Agent 全自主执行，控制栏浮出显示进度条和脉冲指示器。')
bullet('6 种快捷操作：继续写作 / 改写润色 / 提炼摘要 / 生成大纲 / 中文优化 / 扩展内容。')

h3('3.1.3 四编辑器桥接（EditorBridge）')
para('EditorBridge 是连接 AI Agent 与四种编辑器的统一抽象层，采用全局注册表模式（Map<EditorType, EditorAPI>）：')

tbl(['编辑器类型', '桥接状态', '支持操作'], [
    ['文档编辑器（TipTap）', '完全接入', 'append/insert/replace/rewrite/delete + 光标定位 + 选区操作'],
    ['PPT/Slides', '已注册', '追加/替换幻灯片内容、查找替换'],
    ['白板（Whiteboard）', '已注册', '插入文本元素、查找替换'],
    ['思维导图（MindMap）', '已注册', '插入节点、查找替换'],
])

para('每种编辑器在挂载时调用 editorBridge.register(type, api) 注册自身的 EditorAPI 实现。AI Agent 通过 getEditorBridge() 获取当前活跃编辑器的桥接接口，无需感知底层编辑器的具体实现差异。')

page_break()

h2('3.2 富文本编辑器')
para('编辑器基于 TipTap 框架构建，针对中文写作场景进行了深度定制，注册了 18+ 个扩展：')

h3('3.2.1 格式能力')
bullet('段落样式：正文、六级标题（H1-H6），通过工具栏下拉菜单或 Slash 命令快速切换。')
bullet('文字格式：加粗、斜体、下划线、删除线、上标、下标、多色高亮、字体颜色、背景色。')
bullet('段落格式：左对齐、居中、右对齐、两端对齐、行间距调节。')
bullet('列表：有序列表、无序列表、任务清单（支持多级嵌套）。')
bullet('块级元素：引用块、代码块（含语法高亮）、分割线。')
bullet('嵌入元素：链接、图片（本地/URL）、表格（支持行列增删、合并/拆分单元格）、视频嵌入。')

h3('3.2.2 交互功能')
bullet('Slash 命令：任意空行输入 / 调出快捷命令面板，无需鼠标即可插入元素。')
bullet('浮动工具栏：选中文本后自动出现格式快捷工具栏。')
bullet('查找与替换：浮层式 FindReplaceBar，支持正则、全词匹配、统计命中数量、全部替换。')

h2('3.3 文档组织与管理')
tbl(['模块', '主要特性'], [
    ['工作区（Workspace）', '按项目或主题组织文档，支持多工作区创建与切换'],
    ['文档树', '文件夹层级结构，支持文档拖拽移动'],
    ['标签系统', '为文档添加任意标签，跨工作区快速筛选'],
    ['全局搜索', '跨文档全文内容搜索（Ctrl+K 唤起），搜索结果实时更新'],
    ['写作统计与文档关系图', 'WritingStatsView 展示累计字数、文档数量、写作频率等统计；DocumentGraphView 以节点图展示文档之间的引用和关联关系'],
])

h2('3.4 版本历史')
para('系统自动为每篇文档维护版本快照，默认每 5 分钟自动创建存储点。用户在 VersionHistory 组件中可浏览、预览和回溯到任意历史版本。所有版本数据存储在本地 better-sqlite3 数据库中，不依赖网络。支持手动创建命名快照、比较两个版本之间的差异。')

h2('3.5 模板库')
para('内置覆盖学术论文、商业报告、项目提案、周报、演讲稿等常见场景的专业文档模板。TemplatesView 组件以卡片网格展示，按类别筛选。用户可将自己的文档另存为模板供复用。')

h2('3.6 文献管理')
para('轻量级文献管理器支持手动录入参考文献元数据（标题、作者、年份、期刊、DOI 等）。支持 DOI 自动查询、URL 元数据解析、BibTeX 批量导入。支持生成 APA、MLA、GB/T 7714 三种学术引用格式。')

h2('3.7 导出功能')
bullet('Word (.docx)：兼容 Microsoft Word，保留基本样式，适合协作与提交。')
bullet('PDF：支持 A4/A3/Letter 页面尺寸，提供明亮/典雅/暗黑三种排版主题。')
bullet('Markdown (.md)：HTML 内容转换为标准 Markdown 格式。')
bullet('纯文本 (.txt)：去除所有格式，导出干净的文字内容。')
bullet('HTML：带完整样式的独立网页文件。')

h2('3.8 协作功能')
para('启文在保持本地优先理念的同时，引入了 Y.js CRDT 技术，实现了零冲突的多光标实时协作：')
bullet('多光标协作：同一文档可由多位用户同时编辑，每人的光标位置以不同颜色标识。')
bullet('评论系统：CommentPanel 组件支持在文档任意位置添加评论，评论线程展示，支持回复和解决。')
bullet('离线协作：各协作者的本地编辑在重新联网后自动合并，由 Y.js 的 CRDT 算法保证最终一致性。')

page_break()
# ═══════════════════════════════════════════
# 四、数据架构
# ═══════════════════════════════════════════
h1('四、数据架构')

h2('4.1 本地存储引擎')
para('启文 v2.0 采用 better-sqlite3 作为本地数据库引擎，比 v1.x 使用的 sql.js（WebAssembly SQLite）提供了约 5-10 倍的读写性能提升。数据库文件存储于用户应用数据目录：')
bullet('Windows: %APPDATA%\\启文\\data\\qiwen.db')
bullet('macOS: ~/Library/Application Support/启文/data/qiwen.db')
bullet('Linux: ~/.config/启文/data/qiwen.db')

h2('4.2 核心数据表')
tbl(['表名', '主要字段', '说明'], [
    ['workspaces', 'id, name, icon, color, profession', '工作区元信息'],
    ['documents', 'id, title, workspace_id, word_count, updated_at', '文档元信息（不含正文）'],
    ['document_contents', 'document_id, content, updated_at', '文档正文（HTML 格式）'],
    ['document_versions', 'id, document_id, content, created_at', '版本历史快照，每 5 分钟自动创建'],
    ['paper_references', 'id, workspace_id, title, authors, doi, year', '文献库数据'],
    ['app_settings', 'key, value', '应用设置键值对（JSON 序列化）'],
    ['user_profile', 'id, name, is_local, plan', '用户信息'],
    ['ai_conversations', 'id, workspace_id, title', 'AI 对话会话'],
    ['ai_messages', 'id, conversation_id, role, content', 'AI 对话消息记录'],
    ['panel_layouts', 'document_id, layout_config, updated_at', '每个文档独立记住的面板布局'],
    ['sync_metadata', 'table_name, last_sync_at, version', '云同步元数据'],
])

h2('4.3 云同步架构')
para('启文采用 Supabase（基于 PostgreSQL）作为云端数据存储，实现本地与云端的双向同步。同步引擎（syncEngine.ts）采用以下策略：')
bullet('双写机制：写操作先写本地数据库（立即返回），再通过 Supabase REST API 异步写入云端。')
bullet('增量同步：启动时拉取自上次同步时间戳以来的云端变更，仅传输增量数据。')
bullet('离线队列：离线期间的写操作进入 pendingQueue（持久化至 localStorage），联网后按时间顺序批量同步。')
bullet('冲突解决：以 updated_at 时间戳较新者为准。Y.js CRDT 用于协作编辑场景的细粒度冲突合并。')
bullet('数据安全：Supabase RLS 策略确保用户只能访问自己的数据；传输层使用 TLS 1.3 加密。')

page_break()
# ═══════════════════════════════════════════
# 五、插件系统
# ═══════════════════════════════════════════
h1('五、插件系统')
para('启文的插件系统实现了职业化自动推荐与全量离线打包两大创新特性。所有插件代码在编译阶段完整打包于安装包中（pluginRegistry.ts 静态注册表），无需网络即可安装使用。')

h2('5.1 职业-插件推荐机制')
para('系统内置 PROFESSION_PLUGIN_MAP 职业-插件映射注册表，定义了 6 类职业与 17+ 个内置插件之间的对应关系。首次启动时用户在引导页（OnboardingPage）选择职业类型，系统自动检索映射表并激活对应插件组合：')

tbl(['职业类型', '自动推荐插件', '核心特色'], [
    ['学术研究', '引用格式生成 / 论文大纲助手 / 关键词提取', 'APA/MLA/GB/T7714 一键生成'],
    ['法律工作', '条款模板库 / 术语风险检查 / 案件时间线', '500+ 内置条款，模糊表述自动标记'],
    ['教育培训', '教案规划器 / 题目生成器 / 思维导图', '文档内容自动生成测验题目'],
    ['医疗健康', '病历模板库 / 药品速查 / ICD 编码查询', 'ICD-10/11 双语本地查询'],
    ['内容创作', '可读性分析 / 人物关系图 / 文风检测', '高频词汇与被动语态分布分析'],
    ['通用知识', '番茄专注 / 快速便签 / 关键词提取', '25 分钟番茄工作法计时器'],
])

h2('5.2 RDM 科研数据管理插件')
para('RDM（Research Data Management）插件是启文面向实验室科研场景的重量级扩展插件，包含 8 个子模块，具备独立的数据库 Schema、Redux Slice 和 UI 组件：')

tbl(['子模块', '功能'], [
    ['仪表盘（Dashboard）', '实验数据统计、最近活动、项目概览'],
    ['项目（Projects）', '科研项目的创建、跟踪与管理'],
    ['电子实验记录（ELN）', '符合 GLP 标准的电子实验记录'],
    ['库存管理（Inventory）', '试剂、耗材与设备库存管理'],
    ['仪器管理（Instruments）', '仪器预约、使用记录与维护'],
    ['审计追踪（Audit）', '数据变更的完整审计日志'],
    ['审批流程（Approvals）', '实验方案、采购等审批工作流'],
    ['报告生成（Reports）', '自动生成标准格式实验报告'],
])

h2('5.3 插件 API 服务层')
para('插件通过统一的 API 服务层与后端数据交互。API 服务按功能域划分：')
bullet('doiApi.ts：DOI 文献元数据查询，调用 Crossref API 获取论文标题、作者、年份等信息。')
bullet('drugApi.ts：药品信息查询，支持中文药品名与英文通用名的双向检索。')
bullet('icdApi.ts：ICD-10/11 疾病编码本地双语数据库查询。')
bullet('legalApi.ts：法律条款模板库查询，支持按关键词、法律领域等多维度检索。')
bullet('readabilityApi.ts：文本可读性分析，基于中文 NLP 计算平均句长、被动语态比例等。')
bullet('semanticApi.ts：语义分析接口，提供关键词提取、文档聚类和概念关联等功能。')

page_break()
# ═══════════════════════════════════════════
# 六、保存与可靠性
# ═══════════════════════════════════════════
h1('六、保存与可靠性机制')
para('启文设计了多层级的数据保护体系，从根本上解决了桌面写作软件"关闭窗口导致数据丢失"的核心痛点。')

h2('6.1 关闭前握手保存协议')
para('针对 Electron 应用中 window.beforeunload 事件无法执行异步 Promise 操作的技术限制，本发明实现了主进程-渲染进程双向握手保存协议：')
bullet('（1）拦截阶段：主进程监听 mainWindow.on("close") 事件，调用 e.preventDefault() 阻止窗口立即关闭。')
bullet('（2）通知阶段：主进程通过 mainWindow.webContents.send("app-before-close") 向渲染进程发送关闭前通知。')
bullet('（3）写入阶段：渲染进程收到通知后，执行 await autoSave.flushAll()，等待所有 pending 文档的异步 IPC 写入全部完成。')
bullet('（4）确认阶段：渲染进程通过 ipcRenderer.send("flush-complete") 向主进程发送保存完成信号，主进程设置 isReallyClosing = true 后执行真正的关闭。同时设置 3 秒超时兜底。')

h2('6.2 多层自动保存')
bullet('定时自动保存：编辑器内容变更后 600ms 触发 autoSave.schedule()，通过防抖策略避免高频写库。')
bullet('失焦即时保存：编辑器失去焦点或用户 Alt+Tab 切换窗口时，立即触发 autoSave.flush() 跳过计时器。')
bullet('数据库写盘确认：每次 IPC 写入完成后调用 saveDatabase() 将至磁盘文件。')

h2('6.3 启动恢复机制')
para('应用启动时依次尝试：Token 刷新（Supabase auth.refreshSession）→ 本地账号恢复（localStorage + SQLite user_profile 表）→ 本地游客模式，三级降级策略确保任意场景下均可正常进入应用。')

page_break()
# ═══════════════════════════════════════════
# 七、国际化与辅助功能
# ═══════════════════════════════════════════
h1('七、国际化与辅助功能')

h2('7.1 多语言支持')
para('启文通过 i18n/index.ts 实现了中英文界面切换。语言包按模块划分：i18n/zh.ts（中文）和 i18n/en.ts（英文），覆盖所有 UI 文本、提示信息和插件描述。')

h2('7.2 辅助工具')
bullet('写作统计（WritingStatsView）：以可视化图表展示累计写作字数、每日写作趋势、文档数量分布等统计信息。')
bullet('文档关系图（DocumentGraphView）：以力导向图展示文档之间的交叉引用和关联网络。')
bullet('代码查看器（CodeViewer/CodeViewerPage）：支持语法高亮的代码阅读模式，含行号和全屏模式。')
bullet('命令面板（CommandPalette）：Ctrl+P 唤起全局命令搜索，支持模糊匹配。')
bullet('组织管理（OrgManageView）：支持创建和管理组织（团队），分配成员角色和权限。')
bullet('许可证管理（LicenseView）：展示当前授权状态、有效期和功能范围。')

page_break()
# ═══════════════════════════════════════════
# 八、安全与隐私
# ═══════════════════════════════════════════
h1('八、安全与隐私')
bullet('本地优先：所有文档默认仅存储于用户本地设备，不经过任何第三方服务器。')
bullet('传输加密：云同步通道使用 TLS 1.3 + Supabase RLS，每个数据库操作均经用户身份验证。')
bullet('AI 数据隔离：内置豆包 API Key 由主进程代理转发，渲染进程不持有 API Key。')
bullet('离线模式：不联网状态下应用 100% 核心功能可用。')
bullet('沙箱安全：渲染进程以 nodeIntegration: false + contextIsolation: true 运行。')
bullet('数据导出自由：支持导出 Markdown、HTML、纯文本等开放格式。')

page_break()
# ═══════════════════════════════════════════
# 九、状态管理架构
# ═══════════════════════════════════════════
h1('九、状态管理架构')
para('所有全局状态通过 Redux Toolkit 管理，部分 Slice 通过 redux-persist 持久化到 localStorage：')

tbl(['Slice', '持久化', '主要状态'], [
    ['appSlice', '✓', '视图、标签页、面板状态、通知'],
    ['documentsSlice', '✗', '已打开文档内容（内存中）、保存状态'],
    ['workspacesSlice', '✗', '工作区列表（每次从 DB 读取）'],
    ['pluginsSlice', '✓（全量）', '已安装插件列表、启用状态'],
    ['settingsSlice', '✓（全量）', '主题、字体、行高等偏好设置'],
    ['editorSlice', '✗', '字数统计、光标位置、查找面板状态'],
    ['authSlice', '✗', '登录态、本地模式、用户信息'],
    ['referencesSlice', '✗', '文献列表（每次从 DB 读取）'],
    ['panelLayoutSlice', '✓（全量）', '面板布局配置'],
    ['rdmSlice', '✓（全量）', 'RDM 科研数据管理状态'],
])

page_break()
# ═══════════════════════════════════════════
# 十、总结
# ═══════════════════════════════════════════
h1('十、总结')
para('启文 v2.0 代表了 BitWool 数字毛织在桌面写作软件领域的重大技术突破。通过 AI Agent 系统、Panel Grid 可拖拽面板布局、better-sqlite3 本地数据库、Supabase 云同步与 Y.js 实时协作、以及职业化插件自动推荐等核心创新，启文从"带 AI 辅助的写作工具"演进为"AI 主导的智能写作平台"，为深度创作者提供了一套完整的、本地优先的、AI 原生的专业写作解决方案。')

page_break()
# ═══════════════════════════════════════════
# 十一、软件界面截图
# ═══════════════════════════════════════════
h1('十一、软件界面截图')
para('以下为启文 v2.0.0 主要功能界面的运行截图。')

# 截图目录
photo_dir = '/sessions/sleepy-exciting-ritchie/mnt/qiwen_text--photo'

screenshots = [
    ('11.1', '主界面 — 编辑器 + AI 对话面板', 'Panel Grid 分屏布局，左侧文档编辑器 + 右侧 ChatPanel AI 对话面板'),
    ('11.2', 'AI Agent 对话面板', 'ChatPanel 特写，展示流式 SSE 回复、快捷操作按钮、思考过程折叠块'),
    ('11.3', 'AI Action 标签 — Diff 确认卡片', 'ActionConfirm 组件弹出状态，展示 before/after Diff 对比'),
    ('11.4', '工作区 + 文档树', '左侧栏文档树，展示文件夹层级、标签筛选和全局搜索'),
    ('11.5', '插件面板', '右侧面板中的插件市场/已安装插件，展示职业分类和插件卡片'),
    ('11.6', '版本历史', 'VersionHistory 组件，展示时间线式的版本快照列表和预览'),
    ('11.7', '导出对话框', 'ExportDialog 组件，展示 Word/PDF/Markdown/纯文本/HTML 五种格式选择'),
    ('11.8', '设置面板', 'SettingsView 设置面板，展示主题/字体/插件/账户等设置项'),
    ('11.9', '面板布局 — 浮动窗口', 'ChatPanel 脱离网格后的独立浮动窗口效果'),
]

from docx.shared import Inches

for num, title, desc in screenshots:
    h2(f'{num} {title}')
    para(desc)
    img_path = os.path.join(photo_dir, f'{num}.png')
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        # small caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f'图 {num} — {title}')
        r.font.name = 'Microsoft YaHei'
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'[ 图片缺失：{num}.png ]')
        r.font.name = 'Microsoft YaHei'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        r.italic = True


# --- Save ---
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '启文_技术文档_v2.0.1.docx')
doc.save(output_path)
print('Written:', output_path)
